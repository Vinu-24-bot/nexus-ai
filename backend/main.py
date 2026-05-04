"""
BATS ForgePro Interview Evaluator - Enterprise FastAPI Backend (Low-Memory Edition)
"""

import os
import re
import json
import shutil
import hashlib
import subprocess
import httpx
import gc
import asyncio
from datetime import datetime, timedelta
from pathlib import Path
from typing import List
from concurrent.futures import ThreadPoolExecutor

from dotenv import load_dotenv
from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, BackgroundTasks, Request, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from sqlalchemy import text  

# 🚀 UPGRADE: Imported SessionLocal to safely handle background database connections
from database import Base, engine, get_db, SessionLocal
from models import Evaluation, InterviewSession, CandidateFeedback
from schemas import (
    EvaluationRequest, EvaluationResponse, QuestionGenerationRequest,
    JDGenerationRequest, AcknowledgmentRequest, SelectionStatusRequest,
    ResumeUploadResponse, SessionCreateRequest, FeedbackRequest
)
from ai_service import (
    evaluate_candidate, generate_interview_questions,
    generate_jd, get_answer_acknowledgment, transcribe_audio, parse_resume_to_json,
    generate_speech_audio
)

load_dotenv()

Base.metadata.create_all(bind=engine)

UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "./uploads"))
UPLOAD_DIR.mkdir(exist_ok=True)
RECORDINGS_DIR = UPLOAD_DIR / "recordings"
RECORDINGS_DIR.mkdir(exist_ok=True)
RESULTS_DIR = UPLOAD_DIR / "results"
RESULTS_DIR.mkdir(exist_ok=True)
RESUMES_DIR = UPLOAD_DIR / "resumes"
RESUMES_DIR.mkdir(exist_ok=True)

app = FastAPI(
    title="BATS ForgePro Backend",
    description="Enterprise AI-powered candidate evaluation backend",
    version="3.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=False, 
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Accept-Ranges", "Content-Range", "Content-Length", "Content-Type"]
)

app.mount("/uploads", StaticFiles(directory=str(UPLOAD_DIR)), name="uploads")

def generate_candidate_id(name: str, position: str) -> str:
    first_name = re.sub(r'[^a-zA-Z]', '', name.split()[0] if name.strip() else "Unknown")
    role_short = re.sub(r'[^a-zA-Z]', '', position.replace(" ", ""))[:20]
    unique_hash = hashlib.md5(f"{name}{position}{datetime.now().isoformat()}".encode()).hexdigest()[:6]
    return f"BATS-{first_name}_{role_short}-{unique_hash}"

def db_to_response(ev: Evaluation) -> dict:
    return {
        "id": ev.id,
        "candidateName": ev.candidate_name,
        "position": ev.position,
        "date": ev.date,
        "candidate_overview": ev.candidate_overview or "",
        "scores": ev.scores or {"technical_proficiency": 0, "relevance_to_jd": 0, "communication": 0, "confidence_level": 0, "overall_score": 0},
        "sentiment": ev.sentiment or {"rating": "Neutral", "explanation": ""},
        "candidate_status": ev.candidate_status or {"level": "Moderate Confidence", "description": ""},
        "selection_status": ev.selection_status or "pending",
        "strengths": ev.strengths or [],
        "red_flags_or_weaknesses": ev.red_flags_or_weaknesses or [],
        "dynamic_follow_up_questions": ev.dynamic_follow_up_questions or [],
        "hiring_recommendation": ev.hiring_recommendation or "Reject",
        "justification": ev.justification or "",
        "video_filename": ev.video_filename,
        "remarks": ev.remarks
    }

def save_result_file(eval_id: str, candidate_name: str, result_data: dict):
    safe_name = candidate_name.replace(" ", "_").replace("/", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{eval_id}_{safe_name}_{timestamp}.json"
    filepath = RESULTS_DIR / filename
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(result_data, f, indent=2, ensure_ascii=False)
    return str(filepath)

def extract_audio(video_path: str, audio_path: str) -> bool:
    try:
        command = ["ffmpeg", "-threads", "1", "-i", video_path, "-q:a", "0", "-map", "a", audio_path, "-y"]
        subprocess.run(command, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        return True
    except Exception as e:
        print(f"[BATS] Audio Extraction Error: {e}")
        return False

# 🚀 UPGRADE: Independent Background Task to upload massive videos safely without OOMing the main thread
def process_background_video_upload(eval_id: str, video_path_str: str, filename: str, folder_name: str):
    hf_token = os.getenv("HF_TOKEN")
    hf_repo_id = os.getenv("HF_REPO_ID")
    if not hf_token or not hf_repo_id:
        return

    try:
        from huggingface_hub import HfApi
        api = HfApi()
        repo_path = f"{folder_name}/{filename}"
        
        # Uploads via memory-safe stream
        api.upload_file(
            path_or_fileobj=video_path_str,
            path_in_repo=repo_path,
            repo_id=hf_repo_id,
            repo_type="dataset",
            token=hf_token
        )
        cloud_url = f"https://huggingface.co/datasets/{hf_repo_id}/resolve/main/{repo_path}"

        # Secure DB Update
        db = SessionLocal()
        try:
            ev = db.query(Evaluation).filter(Evaluation.id == eval_id).first()
            if ev:
                ev.video_filename = f"[UPLOADED] {cloud_url}"
                db.commit()
        finally:
            db.close()

        # Vaporize local file to protect Render Disk Space
        if os.path.exists(video_path_str):
            os.remove(video_path_str)
        audio_path = video_path_str + ".mp3"
        if os.path.exists(audio_path):
            os.remove(audio_path)

    except Exception as e:
        print(f"[BATS ForgePro] Background Video Upload Error: {e}")

def append_to_hf_dataset(eval_data: dict, candidate_id: str):
    hf_token = os.getenv("HF_TOKEN")
    hf_repo_id = os.getenv("HF_REPO_ID") 
    if not hf_token or not hf_repo_id: return
    try:
        from huggingface_hub import HfApi
        api = HfApi()
        training_row = {
            "instruction": f"Evaluate this candidate for the role of {eval_data['position']}. JD: {eval_data['job_description']}",
            "input": f"RESUME: {eval_data['resume']}\nTRANSCRIPT: {eval_data['transcript']}",
            "output": json.dumps(eval_data['scores']) 
        }
        tmp_path = RESULTS_DIR / f"{candidate_id}_dataset.json"
        with open(tmp_path, "w") as f: json.dump(training_row, f)
        api.upload_file(path_or_fileobj=str(tmp_path), path_in_repo=f"training_data/{candidate_id}.json", repo_id=hf_repo_id, repo_type="dataset", token=hf_token)
    except Exception as e:
        print(f"[BATS] Background Dataset Sync Error: {e}")

async def send_system_email(to_email: str, subject: str, body: str, reply_to: str = ""):
    gas_url = os.getenv("GOOGLE_SCRIPT_URL")
    if not gas_url: return
    payload = {"to": to_email, "subject": subject, "body": body, "replyTo": reply_to}
    try:
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
            response = await client.post(gas_url, json=payload)
            print(f"[BATS EMAIL LOG] Attempted to send to: {to_email} | Reply-To: {reply_to} | GAS Response: {response.text}")
    except Exception as e:
        print(f"[BATS ForgePro] Async Email Send Error: {e}")

class TTSRequest(BaseModel):
    text: str
    gender: str = "female"

@app.post("/api/tts")
async def text_to_speech(req: TTSRequest):
    try:
        audio_bytes = await generate_speech_audio(req.text, req.gender)
        return Response(content=audio_bytes, media_type="audio/mpeg")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/transcribe-chunk")
async def transcribe_chunk(audio: UploadFile = File(...)):
    try:
        temp_path = RECORDINGS_DIR / f"chunk_{datetime.now().timestamp()}.webm"
        with open(temp_path, "wb") as f:
            f.write(await audio.read())
        text = await transcribe_audio(str(temp_path))
        os.remove(temp_path)
        return {"text": text}
    except Exception: return {"text": ""}

@app.get("/")
async def root(): return {"message": "BATS ForgePro Enterprise Backend is awake and running!"}

@app.get("/api/health")
async def health_check(db: Session = Depends(get_db)):
    try:
        db.execute(text("SELECT 1"))
        return {"status": "online", "database": "connected and awake"}
    except Exception as e: return {"status": "offline", "error": str(e)}

@app.get("/api/force-reset-db")
async def force_reset_db():
    try:
        Base.metadata.drop_all(bind=engine)
        Base.metadata.create_all(bind=engine)
        return {"message": "SUCCESS: Database tables wiped and successfully rebuilt with the newest schema columns!"}
    except Exception as e: return {"error": str(e)}

@app.get("/api/stream/{filename}")
async def stream_video(filename: str, request: Request):
    file_path = RECORDINGS_DIR / filename
    if not file_path.exists(): raise HTTPException(status_code=404, detail="File not found")
    file_size = file_path.stat().st_size
    range_header = request.headers.get("range")
    if range_header:
        match = re.search(r'bytes=(\d+)-(\d*)', range_header)
        byte1 = int(match.group(1)) if match else 0
        byte2 = int(match.group(2)) if match and match.group(2) else file_size - 1
        length = byte2 - byte1 + 1
        def file_iterator(path, start, length, chunk_size=8192 * 4):
            with open(path, "rb") as f:
                f.seek(start)
                remaining = length
                while remaining > 0:
                    chunk = f.read(min(chunk_size, remaining))
                    if not chunk: break
                    remaining -= len(chunk)
                    yield chunk
        headers = {
            "Content-Range": f"bytes {byte1}-{byte2}/{file_size}", "Accept-Ranges": "bytes",
            "Content-Length": str(length), "Content-Type": "video/webm" if str(filename).endswith(".webm") else "video/mp4",
        }
        return StreamingResponse(file_iterator(file_path, byte1, length), status_code=206, headers=headers)
    return FileResponse(file_path, media_type="video/webm")

@app.post("/api/sessions/create")
async def create_interview_session(request: Request, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    try:
        data = await request.json()
        
        dur = int(data.get("duration_minutes", 10))
        voice = str(data.get("voice_type", "female"))
        jd = data.get("job_description", "")
        config_str = f"\n\n[FORGEPRO_CONFIG] {json.dumps({'dur': dur, 'voice': voice})}"
        
        new_session = InterviewSession(
            candidate_name=data.get("candidate_name", "Unknown"), candidate_email=data.get("candidate_email", ""),
            recruiter_email=data.get("recruiter_email", ""), interview_level=data.get("interview_level", "L2 (Mid-Level)"),
            position=data.get("position", "Technical Role"), job_description=jd + config_str,
            resume_text=data.get("resume_text", ""), status="pending"
        )
        
        db.add(new_session)
        db.commit()
        db.refresh(new_session)
        
        frontend_url = os.getenv("FRONTEND_URL", "https://nexus-ai-platform-omega.vercel.app").rstrip('/')
        interview_link = f"{frontend_url}/interview/{new_session.id}"
        talent_name = data.get("talent_associate_name", "The ForgePro Team")
        candidate_body = f"""Hello {new_session.candidate_name},\n\nYou have been invited to a BATS ForgePro Initial Screening for the {new_session.position} role ({new_session.interview_level}).\n\nYour Talent Associate, {talent_name}, has set up a secure interview vault for you.\nPlease ensure you are on a laptop/desktop, as you will be required to share your screen and camera to proceed.\n\nStart Interview: {interview_link}\n\nBest regards,\n{talent_name}\nBATS ForgePro Talent Acquisition"""
        background_tasks.add_task(send_system_email, new_session.candidate_email, f"Interview Invitation: {new_session.position}", candidate_body, new_session.recruiter_email)
        
        if new_session.recruiter_email:
            dashboard_link = f"{frontend_url}/dashboard"
            recruiter_body = f"Session Created for {new_session.candidate_name}.\n\nRole: {new_session.position} ({new_session.interview_level})\nCandidate Email: {new_session.candidate_email}\nCreated By: {talent_name}\n\nYou will receive automated alerts when the candidate begins and completes the assessment.\n\nAccess your command center here: {dashboard_link}"
            background_tasks.add_task(send_system_email, new_session.recruiter_email, f"BATS Tracking: Session Created for {new_session.candidate_name}", recruiter_body)
        return {"message": "Session created and emails queued", "session_id": new_session.id}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Database Crash: {str(e)}")

@app.get("/api/sessions/{session_id}")
async def get_interview_session(session_id: str, db: Session = Depends(get_db)):
    session = db.query(InterviewSession).filter(InterviewSession.id == session_id).first()
    if not session: raise HTTPException(status_code=404, detail="This interview link does not exist.")
    if datetime.utcnow() - session.created_at > timedelta(hours=24): raise HTTPException(status_code=403, detail="This interview link has expired (24-hour limit). Please contact your recruiter.")
    if session.status != "pending": raise HTTPException(status_code=403, detail="This interview session has already been attempted or completed and is now permanently locked.")
    
    jd = session.job_description or ""
    dur = 10
    voice = "female"
    if "[FORGEPRO_CONFIG]" in jd:
        parts = jd.split("[FORGEPRO_CONFIG]")
        jd = parts[0].strip()
        try:
            config = json.loads(parts[1].strip())
            dur = config.get("dur", 10)
            voice = config.get("voice", "female")
        except: pass
        
    return { "id": session.id, "candidate_name": session.candidate_name, "position": session.position, "job_description": jd, "resume_text": session.resume_text, "interview_level": session.interview_level, "duration_minutes": dur, "voice_gender": voice, "status": session.status }

@app.patch("/api/sessions/{session_id}/status")
async def update_session_status(session_id: str, request: Request, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    req_data = await request.json()
    status = req_data.get("status")
    remarks = req_data.get("remarks", "")
    session = db.query(InterviewSession).filter(InterviewSession.id == session_id).first()
    if not session: raise HTTPException(status_code=404, detail="Session not found")
    session.status = status
    db.commit()
    if session.recruiter_email:
        frontend_url = os.getenv("FRONTEND_URL", "https://nexus-ai-platform-omega.vercel.app").rstrip('/')
        dashboard_link = f"{frontend_url}/dashboard"
        if status == "started":
            body = f"Alert: {session.candidate_name} has just joined the interview room and started their secure camera/screen-share.\n\nYou can monitor the active pipeline here: {dashboard_link}"
            background_tasks.add_task(send_system_email, session.recruiter_email, f"🟢 Started: {session.candidate_name}", body)
        elif status == "completed":
            is_cheat = "SECURITY BREACH" in remarks
            is_leave = "left early" in remarks
            if is_cheat:
                subject = f"🚨 SECURITY TERMINATION: {session.candidate_name}"
                body = f"URGENT ALERT: The interview for {session.candidate_name} was automatically TERMINATED by the BATS Anti-Cheat System.\n\nReason: {remarks}\n\nClick the link below to view the incident logs and dashboard:\n{dashboard_link}"
            elif is_leave:
                subject = f"⚠️ INCOMPLETE: {session.candidate_name} Left Early"
                body = f"Update: {session.candidate_name} manually exited the interview before completion.\n\nClick the link below to view their partial results on the dashboard:\n{dashboard_link}"
            else:
                subject = f"✅ Session Concluded: {session.candidate_name}"
                body = f"Update: {session.candidate_name} has successfully submitted their interview session.\n\nBATS ForgePro is finalizing the video processing and technical evaluation. Click the link below to view the results and export the data directly from your dashboard:\n{dashboard_link}"
            background_tasks.add_task(send_system_email, session.recruiter_email, subject, body)
    return {"message": f"Status updated to {status}", "candidate": session.candidate_name}

@app.post("/api/feedback")
async def submit_feedback(req: FeedbackRequest, db: Session = Depends(get_db)):
    feedback = CandidateFeedback(session_id=req.session_id, candidate_name=req.candidate_name, rating=req.rating, comments=req.comments)
    db.add(feedback)
    db.commit()
    return {"message": "Feedback saved successfully"}

@app.post("/api/upload-video")
async def upload_video(video: UploadFile = File(...)):
    if not video.filename: raise HTTPException(400, "No filename provided")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = f"{timestamp}_{video.filename}"
    file_path = RECORDINGS_DIR / safe_name
    with open(file_path, "wb") as f: shutil.copyfileobj(video.file, f)
    gc.collect() 
    return {"filename": safe_name, "path": f"recordings/{safe_name}", "size": file_path.stat().st_size, "timestamp": timestamp}

@app.post("/api/upload-resume", response_model=ResumeUploadResponse)
async def upload_resume(file: UploadFile = File(...)):
    if not file.filename: raise HTTPException(400, "No filename provided")
    allowed_extensions = {".pdf", ".txt", ".doc", ".docx", ".md"}
    ext = Path(file.filename).suffix.lower()
    if ext not in allowed_extensions: raise HTTPException(400, f"Unsupported file type: {ext}")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = f"{timestamp}_{file.filename}"
    file_path = RESUMES_DIR / safe_name
    content = await file.read()
    with open(file_path, "wb") as f: f.write(content)
    extracted_text = ""
    if ext in [".txt", ".md"]: extracted_text = content.decode("utf-8", errors="ignore")
    elif ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            extracted_text = "\n".join([page.get_text() for page in doc])
            doc.close()
        except ImportError: extracted_text = "[PDF text extraction requires PyMuPDF.]"
    elif ext in [".doc", ".docx"]:
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            text_parts = []
            for p in doc.paragraphs:
                if p.text.strip(): text_parts.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text: text_parts.append(row_text)
            extracted_text = "\n".join(text_parts)
        except ImportError: extracted_text = "[DOCX extraction requires python-docx.]"
    del content
    gc.collect() 
    return ResumeUploadResponse(filename=safe_name, path=str(file_path), extracted_text=extracted_text, size=file_path.stat().st_size)

@app.post("/api/generate-questions")
async def generate_questions(req: QuestionGenerationRequest):
    try:
        questions = await generate_interview_questions(req.job_description, req.resume, req.num_questions, req.interview_level)
        return {"questions": questions}
    except Exception as e: raise HTTPException(500, detail=f"Question generation failed: {str(e)}")

@app.post("/api/generate-jd")
async def generate_job_description(req: JDGenerationRequest):
    try:
        jd = await generate_jd(req.position)
        return {"job_description": jd}
    except Exception as e: raise HTTPException(500, detail=f"JD generation failed: {str(e)}")

@app.post("/api/acknowledge-answer")
async def acknowledge_answer(req: AcknowledgmentRequest):
    try: return await get_answer_acknowledgment(req.question, req.answer, req.next_question)
    except Exception: return {"response_text": f"Got it. Let's move on to the next question.", "is_sufficient": True}

@app.post("/api/evaluate")
async def create_evaluation(req: EvaluationRequest, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    try:
        clean_remarks = req.remarks or "Completed normally."
        is_l1 = "L1_TECH_ROUND" in clean_remarks.upper()
        
        # 🚀 UPGRADE: Dynamic Smart Folder Routing
        folder_name = "l1_tech_rounds" if is_l1 else "initial_screenings"
        
        final_transcript = req.transcript
        actual_video_filename = req.video_filename.replace("[UPLOADED]", "").strip() if req.video_filename else None
        final_video_storage = req.video_filename
        
        if actual_video_filename and actual_video_filename != "LIVE_SCREENING":
            video_path = RECORDINGS_DIR / actual_video_filename
            if video_path.exists():
                audio_path = RECORDINGS_DIR / f"{actual_video_filename}.mp3"
                if extract_audio(str(video_path), str(audio_path)):
                    try: 
                        new_transcript = await transcribe_audio(str(audio_path))
                        # Prevents overwriting manual L1 transcripts
                        is_placeholder = "(Pre-recorded" in final_transcript if final_transcript else True
                        if is_placeholder or not final_transcript.strip():
                            final_transcript = new_transcript
                        else:
                            final_transcript = final_transcript + "\n\n[AUTO-TRANSCRIPT EXTRACTED]:\n" + new_transcript
                    except: pass
                
                # 🚀 CRITICAL FIX: We DO NOT block to upload to HF here anymore. Fixes Render OOM/Timeout.

        gc.collect()

        behavior_data = {}
        if "METRICS_PAYLOAD:" in clean_remarks:
            try:
                parts = clean_remarks.split("METRICS_PAYLOAD:")
                clean_remarks = parts[0].strip()
                behavior_data = json.loads(parts[1])
            except: pass

        ai_result = await evaluate_candidate(req.job_description, req.resume, final_transcript, req.position, clean_remarks, behavior_data)
        candidate_id = generate_candidate_id(req.candidate_name, req.position)

        evaluation = Evaluation(
            id=candidate_id, candidate_name=req.candidate_name, position=req.position,
            job_description=req.job_description, resume=req.resume, transcript=final_transcript,
            video_filename=final_video_storage, remarks=clean_remarks, 
            candidate_overview=ai_result.get("candidate_overview", ""), scores=ai_result.get("scores", {}),
            sentiment=ai_result.get("sentiment", {"rating": "Neutral", "explanation": ""}),
            candidate_status=ai_result.get("candidate_status", {"level": "Moderate Confidence", "description": ""}),
            strengths=ai_result.get("strengths", []), red_flags_or_weaknesses=ai_result.get("red_flags_or_weaknesses", []),
            dynamic_follow_up_questions=ai_result.get("dynamic_follow_up_questions", []),
            hiring_recommendation=ai_result.get("hiring_recommendation", ""), justification=ai_result.get("justification", ""),
        )
        
        db.add(evaluation)
        db.commit()
        db.refresh(evaluation)

        response_data = db_to_response(evaluation)
        save_result_file(evaluation.id, evaluation.candidate_name, response_data)
        
        # 🚀 UPGRADE: Delegated massive file I/O to background tasks
        if actual_video_filename and actual_video_filename != "LIVE_SCREENING":
            if video_path.exists():
                background_tasks.add_task(process_background_video_upload, candidate_id, str(video_path), actual_video_filename, folder_name)

        dataset_payload = {
            "position": req.position, "job_description": req.job_description,
            "resume": req.resume, "transcript": final_transcript, "scores": ai_result.get("scores", {})
        }
        background_tasks.add_task(append_to_hf_dataset, dataset_payload, candidate_id)

        gc.collect() 
        return response_data

    except ValueError as e: raise HTTPException(400, detail=str(e))
    except Exception as e: raise HTTPException(500, detail=f"AI evaluation failed: {str(e)}")

@app.get("/api/evaluations")
async def list_evaluations(db: Session = Depends(get_db)):
    evaluations = db.query(Evaluation).order_by(Evaluation.created_at.desc()).all()
    return [db_to_response(ev) for ev in evaluations]

@app.get("/api/evaluations/{eval_id}")
async def get_evaluation(eval_id: str, db: Session = Depends(get_db)):
    ev = db.query(Evaluation).filter(Evaluation.id == eval_id).first()
    if not ev: raise HTTPException(404, "Evaluation not found")
    return db_to_response(ev)

@app.patch("/api/evaluations/{eval_id}/status")
async def update_selection_status(eval_id: str, req: SelectionStatusRequest, db: Session = Depends(get_db)):
    ev = db.query(Evaluation).filter(Evaluation.id == eval_id).first()
    if not ev: raise HTTPException(404, "Evaluation not found")
    ev.selection_status = req.status
    db.commit()
    db.refresh(ev)
    return db_to_response(ev)

@app.delete("/api/evaluations/{eval_id}")
async def delete_evaluation(eval_id: str, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    ev = db.query(Evaluation).filter(Evaluation.id == eval_id).first()
    if not ev: raise HTTPException(404, "Evaluation not found")

    video_filename = ev.video_filename
    candidate_id = ev.id

    db.delete(ev)
    db.commit()

    def cleanup_cloud_files(v_file: str, c_id: str):
        hf_token = os.getenv("HF_TOKEN")
        hf_repo_id = os.getenv("HF_REPO_ID")
        
        if v_file:
            actual_filename = v_file.replace("[UPLOADED]", "").strip()
            if "huggingface.co" in actual_filename:
                actual_filename = actual_filename.split("/")[-1]

            local_vid = RECORDINGS_DIR / actual_filename
            local_audio = RECORDINGS_DIR / f"{actual_filename}.mp3"
            
            if local_vid.exists():
                try: os.remove(local_vid)
                except: pass
            if local_audio.exists():
                try: os.remove(local_audio)
                except: pass

            if hf_token and hf_repo_id:
                try:
                    from huggingface_hub import HfApi
                    api = HfApi()
                    # 🚀 UPGRADE: The Triple-Threat Vacuum. Destroys the file regardless of which folder it was saved in!
                    try: api.delete_file(path_in_repo=f"initial_screenings/{actual_filename}", repo_id=hf_repo_id, token=hf_token, repo_type="dataset")
                    except: pass
                    try: api.delete_file(path_in_repo=f"l1_tech_rounds/{actual_filename}", repo_id=hf_repo_id, token=hf_token, repo_type="dataset")
                    except: pass
                    try: api.delete_file(path_in_repo=f"recordings/{actual_filename}", repo_id=hf_repo_id, token=hf_token, repo_type="dataset")
                    except: pass
                except Exception as e: print(f"HF Delete Error: {e}")

        if hf_token and hf_repo_id:
            try:
                from huggingface_hub import HfApi
                api = HfApi()
                api.delete_file(path_in_repo=f"training_data/{c_id}.json", repo_id=hf_repo_id, token=hf_token, repo_type="dataset")
            except Exception as e: print(f"HF Delete Error: {e}")

    background_tasks.add_task(cleanup_cloud_files, video_filename, candidate_id)
    return {"message": "Candidate deleted successfully and cloud storage cleanup initiated"}

@app.get("/api/stats")
async def get_stats(db: Session = Depends(get_db)):
    evaluations = db.query(Evaluation).all()
    total = len(evaluations)
    if total == 0:
        return {"total": 0, "avg_score": 0, "strong_hires": 0, "lean_hires": 0, "rejects": 0, "selected": 0, "rejected": 0, "pending": 0, "pipeline_health": "No Data", "top_scorer": None, "positions": []}
    scores = [ev.scores.get("overall_score", 0) if ev.scores else 0 for ev in evaluations]
    avg_score = round(sum(scores) / total, 1)
    strong_hires = sum(1 for ev in evaluations if ev.hiring_recommendation == "Strong Hire")
    lean_hires = sum(1 for ev in evaluations if ev.hiring_recommendation == "Lean Hire")
    rejects = sum(1 for ev in evaluations if ev.hiring_recommendation == "Reject")
    
    if total > 0:
        hire_rate = (strong_hires + lean_hires) / total
        pipeline_health = "Excellent" if hire_rate >= 0.4 else "Needs Adjustment" if hire_rate < 0.15 else "Healthy"
    else: pipeline_health = "No Data"
    top_eval = max(evaluations, key=lambda e: (e.scores or {}).get("overall_score", 0))

    return {"total": total, "avg_score": avg_score, "strong_hires": strong_hires, "lean_hires": lean_hires, "rejects": rejects, "selected": sum(1 for ev in evaluations if ev.selection_status == "selected"), "rejected": sum(1 for ev in evaluations if ev.selection_status == "rejected"), "pending": sum(1 for ev in evaluations if ev.selection_status == "pending"), "pipeline_health": pipeline_health, "top_scorer": top_eval.candidate_name, "positions": list(set(ev.position for ev in evaluations))}

@app.get("/api/feedback")
async def get_all_feedback(db: Session = Depends(get_db)):
    feedbacks = db.query(CandidateFeedback).order_by(CandidateFeedback.id.desc()).all()
    return [{"id": f.id, "candidate": f.candidate_name, "rating": f.rating, "comments": f.comments} for f in feedbacks]

@app.delete("/api/feedback/{feedback_id}")
async def delete_feedback(feedback_id: str, db: Session = Depends(get_db)):
    try:
        search_id = int(feedback_id) if feedback_id.isdigit() else feedback_id
        feedback = db.query(CandidateFeedback).filter(CandidateFeedback.id == search_id).first()
        if feedback:
            db.delete(feedback)
            db.commit()
        return {"message": "Feedback permanently erased"}
    except Exception as e: raise HTTPException(status_code=500, detail=f"Database Lock or Error: {str(e)}")